# Personal info we need to extract from resume/user_input to make a cover letter:
# - Name
# - Address
# - Phone Number
# - Email
# - LinkedIn
# - GitHub
# - Portfolio
# - Summary (About Me)
# - Experience
# - Skills
# - Education
# - Awards
### Extracting personal information from resume

from utils import JobToID

#Custom imports
from configs import load_config,get_resume_path,update_config_with_personal_info,load_secrets
# Import the necessary modules
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate
from pydantic import BaseModel,Field
from langchain_openai import ChatOpenAI

from langchain_community.document_loaders import PyPDFLoader
from langchain_core.prompts import ChatPromptTemplate
# from PyPDF2 import PdfReader
from langchain_google_genai import ChatGoogleGenerativeAI

from datetime import datetime
import configs
import os
import getpass
import glob
from pathlib import Path
import json

secrets = load_secrets()
if not os.environ.get("GOOGLE_API_KEY"):
  print("GOOGLE_API_KEY enviroment variable not set")
  os.environ["GOOGLE_API_KEY"] = getpass.getpass("Enter API key for Gemeni Api (Sign up for free here: https://aistudio.google.com/apikey): ")

# Initialize the LLM model with langchain
if "GOOGLE_API_KEY" in os.environ:
    model = ChatGoogleGenerativeAI(model="gemini-2.0-flash-exp", temperature=0, max_tokens=None, timeout=None, max_retries=2)
    print("Using Gemeni AI")
else:
    model = ChatOpenAI(model="gpt-4o-mini")
    print("Using OpenAI")

# ---------- HELPER FUNCTION ----------
# Function to extract personal information from a resume
def get_all_pdf_text(pdf_path: str) -> str:
    """
    Extracts text from all pages of a PDF file.
    
    Parameters:
        pdf_path (str): Path to the PDF file
        
    Returns:
        str: Combined text from all pages of the PDF
    """
    loader = PyPDFLoader(pdf_path)
    pages = loader.load()
    return ' '.join(page.page_content for page in pages)

# ------------- personal information extraction ----------------

# Define your desired data structure.
class PersonalInformation(BaseModel):
    name: str = Field(description="The person's first name")
    surname: str = Field(description="The person's last name")
    email: str = Field(description="The person's email address")
    phone: str = Field(description="The person's phone number")
    linkedin_url: str = Field(description="The person's LinkedIn profile URL")
    github_url: str = Field(description="The person's GitHub profile URL")
    portfolio_url: str = Field(description="The person's personal portfolio URL")
    address: str = Field(description="The person's address")
    summary: str = Field(description="A brief summary about the person")
    experience: list = Field(description="A list of the person's work experiences")
    skills: list = Field(description="A list of the person's skills")
    education: list = Field(description="A list of the person's educational background")

parser = JsonOutputParser(pydantic_object=PersonalInformation)

# Function to extract personal information from a resume
def extract_personal_information_from_resume(resume_text):
    from configs import get_personal_info_from_resume  # Moved inside the function
    # Define the prompt template for extracting personal information from a resume
    prompt = PromptTemplate(
        template="Answer the user query.\n{format_instructions}\n{query}\n",
        input_variables=["Extract all relevant personal information from the resume."],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )

    # Define the chain of operations for extracting personal information
    chain = prompt | model | parser

    # Invoke the chain with the resume text
    try:
        result = chain.invoke({"query": resume_text})
    except Exception as e:
        print(f"Error extracting personal information: {e}")
        result = PersonalInformation(
            name="",
            surname="",
            email="",
            phone="",
            linkedin_url="",
            github_url="",
            portfolio_url="",
            address="",
            summary="",
            experience=[],
            skills=[],
            education=[],
        )
    return result

# ------ extract json from job description ------
class JobInformation(BaseModel):
    title: str = Field(description="The job title")
    company: str = Field(description="The company offering the job")
    location: str = Field(description="The job location")
    responsibilities: list = Field(description="A list of job responsibilities")
    qualifications: list = Field(description="A list of job qualifications")
    benefits: list = Field(description="A list of job benefits")
    skills: list = Field(description="A list of required technical skills for the job, such as SQL, Python, etc.")

def extract_job_information(job_description: str) -> JobInformation:
    """
    Extracts structured job information from a job description string.
    Parameters:
        job_description (str): The job description text to extract information from.
    Returns:
        JobInformation: A data structure containing the extracted job details.
    """

    # Initialize the JSON output parser with the defined data structure
    job_info_parser = JsonOutputParser(pydantic_object=JobInformation)
    
    # Define the prompt template for extracting job information
    job_info_prompt = PromptTemplate(
        template="Answer the user query.\n{format_instructions}\n{query}\n",
        input_variables=["Extract all relevant job information from the job description."],
        partial_variables={"format_instructions": job_info_parser.get_format_instructions()},
    )
    
    # Create a chain of the prompt and model, then parse the response
    chain = job_info_prompt | model | job_info_parser
    
    # Invoke the chain with the job description to extract job information
    try:
        job_info = chain.invoke({"query": job_description})
    except Exception as e:
        print(f"Error extracting job information: {e}")
        job_info = JobInformation(
            title="",
            company="",
            location="",
            responsibilities=[],
            qualifications=[],
            benefits=[],
            skills=[]
        )
    return job_info

    ### ---------- Revised Function for Creating Cover Letter Content ------------
def create_tailored_cover_letter_content(job_description, resume):
    
    cover_letter_prompt_template_e = ChatPromptTemplate([
        ("system", "You are a professional career coach specializing in crafting concise, compelling, and highly tailored cover letters that align a candidate's strengths with a specific job."),   
        ("user", """
        Craft a personalized and engaging cover letter based on the provided job description and resume.

        ## Guidelines:
        *   **Length:** Up to three well-structured paragraphs.
        *   **Tone:** Warm, confident, and enthusiastic while maintaining professionalism. Aim to connect emotionally with the reader.
        *   **Content Breakdown:**
            *   **Introduction:** Begin with an attention-grabbing statement that introduces the candidate, specifies the desired position and company, and conveys genuine excitement about the opportunity. Reference a specific company value, mission, or recent accomplishment that aligns with the candidate's interests.
            *   **Body:**
                *   Highlight 2-3 standout achievements or skills from the resume that are directly relevant to the job requirements. Use concrete, quantifiable examples to demonstrate the candidate’s contributions and impact.
                *   Illustrate how the candidate's professional goals and personal values align with the company's mission, values, or ongoing projects. Showcase a deep understanding of the company through thoughtful connections.
            *   **Conclusion:** Summarize why the candidate is a perfect fit for the role, expressing eagerness to contribute to the company's success. End with a forward-looking statement that invites further discussion or an interview.
        *   **Formatting:** Write in clear and natural language with fluent transitions between paragraphs. Avoid repetitive phrases, technical jargon, and placeholder text (e.g., [Company Name]).Use clear, concise paragraphs. Avoid overly technical jargon or placeholders (e.g., [Company Name]).Use concise paragraphs. Do not include a salutation, closing (e.g., "Sincerely"), or signature.
        *   **Style Enhancements:** 
            * Incorporate storytelling elements to make the narrative compelling and memorable.
            * Use varied sentence structures and strong action verbs to maintain reader interest.
            * Balance technical proficiencies with soft skills to present a well-rounded candidate profile.

        ## Provided Information:
        *   **Job Description:** {job_description}
        *   **Resume:** {resume}

        ## Response:
        Deliver only the cover letter text in a cohesive and engaging manner. Exclude any supplementary explanations, analyses, or placeholder content.

        """)])

    promt_for_optimized_cover_letter = cover_letter_prompt_template_e.format(job_description=job_description, resume=resume)
    optimized_cover_letter = model.invoke(promt_for_optimized_cover_letter)
    #check for errores
    if "[" in optimized_cover_letter.content:
        print("Error: Cover letter may contain placeholders. Please revise the content/prompt to remove placeholders.")
    return optimized_cover_letter.content
# use this personal_info json and the job_information json extracted from the job description to create a customized cover letter

# Create the cover letter from the extracted personal information and job information

# --------------- Document Generation ----------------
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from utils import add_hyperlink

# Lib for docx2pdf
from docx2pdf import convert
# Load personal information from the resume

# --------------- Document Generation ----------------
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from utils import add_hyperlink
from datetime import datetime
# Lib for docx2pdf

from docx2pdf import convert
# Load personal information from the resume

from configs import get_personal_info_from_resume  # Added import

# Replace the undefined function with the imported one
resume_personal_info = get_personal_info_from_resume()
# Load yaml file personal information
personal_info = ()

def generate_tailored_cover_letter(job_description_json, personal_info_json,cover_letter_text):

    # personal_info = load_config()
    personal_info = personal_info_json
    # job_description = job_description_json["job_description"]
    # personal_info = personal_info_json['personal_information']
    full_name = personal_info["name"] + " " + personal_info["surname"]
    linkedin_url = personal_info["linkedin"]
    github_url = personal_info["github"]
    portfolio_url = personal_info["portfolio"]
    email = personal_info["email"]
    phone = personal_info["phone"]
    
    job_title = job_description_json["title"]
    company = job_description_json["company"]
    location = job_description_json["location"]

    # Initialize the document
    doc = Document()

    # Add a centered title to the document with the full name and contact info
    title_info_heading = doc.add_heading("", level=1)
    title_info_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_info_heading.add_run(f"{full_name}\n")

    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
    title_info_heading.add_run(f"{email} | {phone}\n").font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
    
    # Add centered hyperlinks for personal info urls if not empty (github, linkedin, portfolio)
    # Add personal info links if they exist and are not empty
    urls = []
    if github_url:
        urls.append(('Github', github_url))
    if linkedin_url:
        urls.append(('LinkedIn', linkedin_url))
    if portfolio_url:
        urls.append(('Portfolio', portfolio_url))
        
    # Add links separated by ' | ' if they exist
    for i, (text, url) in enumerate(urls):
        if i > 0:
            title_info_heading.add_run(' | ')
        add_hyperlink(paragraph=title_info_heading, text=text, url=url)
    
    # Add today's date and job information as one paragraph with newlines 
    job_info_paragraph = doc.add_paragraph()
    job_info_paragraph.add_run(f"{datetime.today().strftime('%B %d, %Y')}\n")
    job_info_paragraph.add_run(f"{company}\n")
    job_info_paragraph.add_run(f"{job_title}\n")
    job_info_paragraph.add_run(f"{location}\n")

    #Add dear hiring manager salutation as one paragraph with newlines
    salutation_paragraph = doc.add_paragraph()
    salutation_paragraph.add_run("Dear Hiring Manager,\n")

    # Add the cover letter text customized by LLM 
    doc.add_paragraph(cover_letter_text)

    # Add a goodbye signature as one paragraph with newlines
    goodbye_paragraph = doc.add_paragraph()
    goodbye_paragraph.add_run("Sincerely,\n")
    goodbye_paragraph.add_run(f"{full_name}.")

    # function for document saving
    def save_document(doc, output_folder, base_filename, extension):
        #later on this will be used to save the document as a pdf and then sent back to the front end

        # Ensure the output folder exists
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        # Set the filename and ensure it is unique
        filename = base_filename + extension
        # counter = 1

        # # Ensure the filename is unique
        # while os.path.exists(os.path.join(output_folder, filename)):
        #     filename = f"{base_filename}_{counter}{extension}"
        #     counter += 1

        # get absolute path of the saved document
        doc_path = Path(output_folder) / filename
        # Save the document
        doc.save(doc_path)
        print(f"Document saved as: {filename}, to the folder: {output_folder}")
        convert(os.path.join(output_folder, filename), os.path.join(output_folder, base_filename + ".pdf"))
        print(f"Document saved as: {base_filename}.pdf, to the folder: {output_folder}")
        # return absolute path of the saved document
        doc_path = Path(output_folder) / filename
        return str(doc_path.resolve())


    # Save the document
    output_folder = './data_folder/output/tailored_cover_letters/'

    #base filename should be named after the job
    job_name_for_file = JobToID(f'{job_description_json["company"]}_{job_description_json["title"]}_cover_letter')
    base_filename = job_name_for_file
    extension = '.docx'
    filename = base_filename + extension
    counter = 1
    
    
    saved_doc_path = save_document(doc, output_folder, base_filename, extension)
    return saved_doc_path
    # return doc

# Generate the tailored cover letter

# doc_test = generate_tailored_cover_letter(json_job_description, resume_personal_info, optimized_cover_letter_text)

# Save the document
def make_and_save_cv_from_job_desc(job_description_text,personal_info, resume_text):
    # Load personal information from the resume
    # Extract job information from the job description text
    job_description_json = extract_job_information(job_description_text)
    # Create a tailored cover letter
    optimized_cover_letter_text = create_tailored_cover_letter_content(job_description_text, resume_text)
    # Generate the tailored cover letter document
    saved_doc_path = generate_tailored_cover_letter(job_description_json, personal_info, optimized_cover_letter_text)
    return saved_doc_path


#test the funciton
def test_make_and_save_cv():
    configs = load_config()
    configs = update_config_with_personal_info(configs)
    print(configs)
    job_description = """May 2025 Geotechnical/Geological Engineering Co-op Sparwood, BC / Elkford, BC Coal Division - Campus /On-site Elk Valley Resources (EVR) operates four steelmaking coal mines that employ over 5,000 people in the Elk Valley of British Columbia. We are committed to responsible resource development, environmental performance and building strong partnerships with communities and Indigenous Peoples. EVR is proud to be part of the Glencore Group, joining a portfolio of operations around the world and a history that dates back more than 100 years in Canada. Start Date: May 2025 Duration: 8 or 12 Months What we offer: At EVR, we're dedicated to supporting our workforce on their unique career journeys. EVR recognizes our employees' skills, competencies, and performance and offers opportunities with competitive compensation and benefits. We also aim to promote employee health and wellness and provide opportunities for professional development. We believe in the power of diversity and strive each day to create a workplace where each individual is valued and respected for their contributions. We also believe that our employees are the reason we are accomplishing great things and are committed to supporting them through our many different programs. Join us in the breathtaking Elk Valley of British Columbia. Here you will find outdoor adventure at your fingertips. Whether it's biking and skiing, or the laid-back atmosphere of fishing and hiking, there is something for everyone! Come experience what work life balance is all about! In this role, you will: Participate in and adhere to health and safety programs and policies Perform a variety of geotechnical assessments for current and future pit slope excavations and waste rock dumps Perform a variety of stability and rockfall analyses for pit slopes and natural slopes Assist with quarterly geological model updates and reconciliations Conduct mapping and sampling activities in all weather conditions Working with consultants to complete geotechnical design work Support Geotechnical monitoring and analysis including interpretation of ground movement and piezometric data, drilling data centralization, instrumentation maintenance, and document tracking, control and critical geotechnical risk management Provide hazard management support Create and evaluate innovative ideas and technologies to help improve the way we do business Requirements: Working towards an Undergraduate degree in Geological or Geotechnical Engineering A valid Class 5 Driver's License or extra-provincial equivalent, without learner's restrictions Able to operate vehicles at our operations, in various weather conditions and terrain Motivated self-starters with the ability to manage several projects simultaneously Comfortable working in a dynamic work environment with a willingness to work outdoors Strong written, verbal and interpersonal communication skills with a capability to communicate complex ideas to all levels and partners Proficient in Microsoft Office applications (Word, Excel, PowerPoint, Access) Benefits: Receive financial assistance for travel expenses and co-op fees Acquire hands-on experience applying academic knowledge to real-world mining operations, fostering the development of technical skills Benefit from mentorship and guidance offered by seasoned specialists in their respective fields Gain a competitive edge, securing primary consideration for professional positions upon graduation $4,642 - $5,883 a month Salary is commensurate with the number of academic semesters and prior co-op experiences completed. Application Deadline: January 31, 2025"""
    filepath = make_and_save_cv_from_job_desc(job_description,configs,configs["personal_information"])
    print(f"Docx path: {filepath}, PDF path: {filepath.replace('.docx','.pdf')}")
    #get quick user feedback requesting to check the generated cv before continuing
    feedback_flag = input("Please check the generated CV in the link above before continuing. If you are satisfied with the result, please type 'y/yes' to continue and the script will complete ALL your applications in a matter of minutes.")
    if feedback_flag.lower() in ['y','yes']:
        print("Continuing with the script")
    else:
        print("Please check the generated CV, your resume content, or the personal_info.yaml file before continuing")
        raise Exception("Please check the generated CV, your resume content, or the personal_info.yaml file before continuing")
    return filepath


# filepath=test_make_and_save_cv()
# pdf_filepath = filepath.replace('.docx','.pdf')
# print(pdf_filepath)

# def make_cover_letter(job_description_txt, cover_letter_filename='tailored_cover_letter'):
#     """
#     Creates a tailored cover letter based on a job description and personal information.
#     This function orchestrates the process of generating a customized cover letter by:
#     1. Loading personal information from a JSON file
#     2. Extracting relevant information from the job description
#     3. Creating tailored content
#     4. Generating a formatted document
#     5. Saving the document to a specified location
#     Parameters:
#         job_description_txt (str): The job description text to analyze and tailor the cover letter to
#         cover_letter_filename (str, optional): Base filename for the output document. 
#                                              Defaults to 'tailored_cover_letter'
#     Returns:
#         None: The function saves the generated cover letter as a Word document in the
#               '../data_folder/output/tailored_cover_letters/' directory
#     Example:
#         >>> make_cover_letter("Software Developer position at Tech Corp")
#         # Creates 'tailored_cover_letter.docx' in the output directory
#     """
#     personal_info = load_personal_info_from_json()
#     job_description = extract_job_information(job_description_txt)
#     cover_letter_text = create_tailored_cover_letter_content(job_description_txt, personal_info)
#     doc = generate_tailored_cover_letter(job_description, personal_info, cover_letter_text)
#     output_folder = '../data_folder/output/tailored_cover_letters/'
#     base_filename = cover_letter_filename
#     extension = '.docx'
#     save_document(doc, output_folder, base_filename, extension)
#     #return full filepath
#     # Return full file path of pdf
#     output_path = os.path.abspath(os.path.join(output_folder, base_filename + '.pdf'))
#     if os.path.exists(output_path):
#         return output_path
#     else:
#         raise FileNotFoundError(f"PDF file not found at {output_path}")


# # Main
# if __name__ == "__main__":
    
#     # Load personal information from the resume
#     # personal_info = load_personal_info_from_json()
#     # print(f"Personal information: {personal_info}")
    
#     # # read file
#     with open(os.path.join(os.path.dirname(__file__), '..', 'data_folder', 'input', 'waste', 'op_tech_coop_description_sample.txt'), 'r', encoding='utf-8') as file:  job_description = file.read()
#     pdf_cover_letter_path = make_cover_letter(job_description, 'custom_cover_letter')
#     # json_job_description = extract_job_information(job_description)
#     # print(f"Job description: {json_job_description}")
#     # output_folder = '../data_folder/output/tailored_cover_letters/'
#     # base_filename = 'custom_cover_letter'
#     # extension = '.docx'

#     # # Create a tailored cover letter
#     # cover_letter_text = create_tailored_cover_letter_content(job_description, personal_info)
#     # doc = generate_tailored_cover_letter(json_job_description, personal_info, cover_letter_text)
#     # save_document(doc, output_folder, base_filename, extension)


